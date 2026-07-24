# extractor.py

import os
import json
import pdfplumber
import pandas as pd
from docx import Document
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM

from extractor_template import (
    IOTAgreement,
    AgmtHeaderStg,
    AgmtModelsStg,
    AgmtMdlNormalStg,
    AgmtCommitment,
)


class Extractor:

    def __init__(self):
        model = "meta-llama/Llama-3.2-3B-Instruct"
        self.tokenizer = AutoTokenizer.from_pretrained(model)
        self.model = AutoModelForCausalLM.from_pretrained(
            model, torch_dtype=torch.float16, device_map="cuda"
        )
        self.info = None

    def read(self, path):
        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            with pdfplumber.open(path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf)

        if ext == ".docx":
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)

            for t in doc.tables:
                text += "\n" + "\n".join(
                    " | ".join(c.text for c in r.cells) for r in t.rows
                )
            return text

        if ext in [".xlsx", ".xls"]:
            sheets = pd.read_excel(path, sheet_name=None)
            return "\n".join(f"{name}\n{df}" for name, df in sheets.items())

        raise ValueError("Unsupported file type")

    def extract(self, text):
        # The prompt is enhanced to explicitly show the LLM the required variables from your schemas.
        prompt = f"""
Extract telecom agreement data from the document below.
Return ONLY valid JSON. Do not include markdown formatting, explanations, or code blocks.

Schema:
{{
  "header": {{
    "agmt_id": "string", "sender": "string", "rp": "string", "tap_direction": "string", 
    "rev_no": 0, "start_date": "string", "end_date": "string", "remarks": "string", 
    "data_level": "string", "currency_code": "string", "agmt_status": "string", 
    "agmt_type": "string", "is_tap_level_agmt": false
  }},
  "models": [
    {{
      "model_seq": 0, "agmt_id": "string", "model_type": "string", "model_name": "string"
    }}
  ],
  "normal_models": [
    {{
      "agmt_id": "string", "model_seq": 0, "rec_type": "string", "zone_code": "string", 
      "rate_currency": "string", "pra_rate_type": "string", "disc_rate_perc": 0.0, 
      "charge_include_tax": false, "charge_field": "string"
    }}
  ],
  "commitments": [
    {{
      "agmt_id": "string", "commitment_name": "string", "commitment_type": "string", 
      "direction": "string", "amount": 0.0, "capture_rate_pct": 0.0, 
      "party_from": "string", "party_to": "string"
    }}
  ]
}}

Document:
{text}
"""

        messages = [{"role": "user", "content": prompt}]
        input_ids = self.tokenizer.apply_chat_template(
            messages, add_generation_prompt=True, return_tensors="pt"
        ).to("cuda")

        result = self.model.generate(input_ids, max_new_tokens=3000, do_sample=False)

        # IMPORTANT: only decode the newly generated tokens.
        # result[0] contains prompt tokens + generated tokens concatenated.
        # Decoding the whole thing and then doing output.find("{") finds the
        # "{" inside the Schema block of the PROMPT itself (the schema is
        # literal JSON with placeholder values like "string"/0/false), which
        # is exactly why you were getting the schema echoed back verbatim.
        new_tokens = result[0][input_ids.shape[-1] :]
        output = self.tokenizer.decode(new_tokens, skip_special_tokens=True)

        data = self._parse_json(output)

        # Fill the Pydantic instances mapping directly to your extractor_template definitions
        self.info = IOTAgreement(
            header=AgmtHeaderStg(**data.get("header", {})),
            models=[AgmtModelsStg(**x) for x in data.get("models", [])],
            normal_models=[
                AgmtMdlNormalStg(**x) for x in data.get("normal_models", [])
            ],
            commitments=[AgmtCommitment(**x) for x in data.get("commitments", [])],
        )

        return self.info

    def _parse_json(self, output: str) -> dict:
        """
        Parses the model's raw text output into a dict, tolerating the
        common ways small instruct models mangle JSON:
          - wrapping the JSON in ```json ... ``` fences
          - trailing commas before a closing } or ]
          - generation cutting off mid-object (missing closing braces)
        Raises ValueError with the raw output attached for debugging if
        nothing works.
        """
        text = output.strip()

        # Strip markdown code fences if present
        if text.startswith("```"):
            text = (
                text.split("```", 2)[1] if text.count("```") >= 2 else text.lstrip("`")
            )
            text = text.strip()
            if text.startswith("json"):
                text = text[4:].strip()

        start_idx = text.find("{")
        end_idx = text.rfind("}") + 1
        if start_idx == -1:
            raise ValueError(
                f"Failed to generate valid JSON (no '{{' found).\nRaw model output:\n{output}"
            )

        candidate = text[start_idx:end_idx] if end_idx > start_idx else text[start_idx:]

        # First attempt: parse as-is
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass

        # Second attempt: remove trailing commas before } or ]
        import re

        repaired = re.sub(r",\s*([}\]])", r"\1", candidate)
        try:
            return json.loads(repaired)
        except json.JSONDecodeError:
            pass

        # Third attempt: output was likely truncated mid-generation
        # (hit max_new_tokens before the model finished). Try closing
        # any unbalanced braces/brackets.
        open_braces = repaired.count("{") - repaired.count("}")
        open_brackets = repaired.count("[") - repaired.count("]")
        patched = repaired.rstrip()
        # Drop a dangling partial key/value fragment ending mid-string or mid-comma
        patched = re.sub(r",\s*\"?[^\"\{\}\[\]]*$", "", patched)
        patched += "]" * max(open_brackets, 0) + "}" * max(open_braces, 0)
        try:
            return json.loads(patched)
        except json.JSONDecodeError as e:
            raise ValueError(
                "Failed to parse model output as JSON after repair attempts.\n"
                f"Parse error: {e}\n"
                f"Raw model output:\n{output}"
            )

    def get_info(self):
        return self.info


def get_contents(file_path: str):
    """
    Opens the document at file_path, reads and understands the text,
    and fills out the IOTAgreement template components.

    Returns a tuple: (header, models, normal_models, commitments)
    """
    extractor = Extractor()
    text = extractor.read(file_path)

    # Process text through the LLM and construct the Pydantic objects
    agreement_data = extractor.extract(text)

    # Return the unpacked attributes as requested
    return (
        agreement_data.header,
        agreement_data.models,
        agreement_data.normal_models,
        agreement_data.commitments,
    )
