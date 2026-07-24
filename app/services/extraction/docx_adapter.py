"""
Word document (.docx) extraction adapter.

Extracts tables and text from .docx files using python-docx.
Consumed by extraction_agent.py for Word document processing.

Includes detection for key-value tables (2-column definition/metadata tables)
and row-length normalization for merged cells.

Dependency: python-docx>=1.1.2

Path: app/services/extraction/docx_adapter.py
      (Replaces existing empty file — was declared but never implemented.)
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Threshold for key-value table detection.
# If average length of cells in the second column exceeds this,
# the table is likely a key-value/definitions table, not a data table.
_KV_AVG_LENGTH_THRESHOLD = 80

# Minimum ratio of rows where column 2 is longer than column 1
# for the table to be classified as key-value.
_KV_RATIO_THRESHOLD = 0.6


def extract_tables_from_docx(file_bytes: bytes) -> list[dict]:
    """
    Extract all tables from a .docx file.

    Handles two table formats:
    - Standard data tables: row 0 is headers, remaining rows are data.
    - Key-value tables: 2-column tables where row 0 is data, not a header.
      Detected heuristically and returned with headers ["Field", "Value"].

    All rows are normalized to consistent length to handle merged cells.

    Args:
        file_bytes: Raw .docx file content as bytes.

    Returns:
        List of dicts, each containing:
            "title"   - Label string (e.g. "Table 1" or "Table 3 (key-value)")
            "headers" - List of column header strings
            "rows"    - List of row data (each row is a list of strings)
        Returns empty list if no tables found or on failure.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return []

    tables_found = []

    try:
        doc = Document(io.BytesIO(file_bytes))

        for table_idx, table in enumerate(doc.tables):
            rows_data = []

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)

            # Skip completely empty tables
            if not rows_data:
                continue

            # Normalize row lengths — merged cells can produce ragged rows
            rows_data = _normalize_row_lengths(rows_data)

            # Skip single-row tables (no usable data)
            if len(rows_data) < 2:
                continue

            # Filter out completely empty rows
            rows_data = [row for row in rows_data if any(cell for cell in row)]

            if len(rows_data) < 2:
                continue

            # Detect key-value tables
            if _is_key_value_table(rows_data):
                # All rows are data, no header row
                tables_found.append(
                    {
                        "title": f"Table {table_idx + 1} (key-value)",
                        "headers": ["Field", "Value"],
                        "rows": rows_data,
                    }
                )
            else:
                # Standard table: row 0 = headers, rest = data
                headers = rows_data[0]
                data_rows = [row for row in rows_data[1:] if any(cell for cell in row)]

                if data_rows:
                    tables_found.append(
                        {
                            "title": f"Table {table_idx + 1}",
                            "headers": headers,
                            "rows": data_rows,
                        }
                    )

    except Exception as e:
        logger.error(f"Word table extraction failed: {e}")
        return []

    logger.info(f"Extracted {len(tables_found)} tables from .docx")
    return tables_found


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all paragraph text from a .docx file.

    Concatenates all non-empty paragraphs with newlines.
    Used as supplementary context for Gemini when tables alone
    don't capture all structured data (e.g. key-value pairs in prose).

    Args:
        file_bytes: Raw .docx file content as bytes.

    Returns:
        Full text content of the document, or "" on failure.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed.")
        return ""

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        logger.info(f"Extracted {len(text)} chars from .docx paragraphs")
        return text
    except Exception as e:
        logger.error(f"Word text extraction failed: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════════════
# Internal helpers
# ═══════════════════════════════════════════════════════════════════════


def _normalize_row_lengths(rows: list[list[str]]) -> list[list[str]]:
    """
    Ensure all rows have the same number of columns.

    Merged cells in Word tables produce rows with fewer cells than
    the table's actual column count. This pads short rows with empty
    strings and truncates overly long rows to the most common width.

    Args:
        rows: Raw extracted rows (possibly ragged).

    Returns:
        List of rows, all with consistent column count.
    """
    if not rows:
        return rows

    # Determine target width as the most common row length.
    # Using mode instead of max avoids inflating width due to
    # a single malformed row with duplicated merged-cell content.
    length_counts: dict[int, int] = {}
    for row in rows:
        length = len(row)
        length_counts[length] = length_counts.get(length, 0) + 1

    target_width = max(length_counts, key=length_counts.get)

    normalized = []
    for row in rows:
        if len(row) < target_width:
            # Pad short rows with empty strings
            row = row + [""] * (target_width - len(row))
        elif len(row) > target_width:
            # Truncate overly long rows
            row = row[:target_width]
        normalized.append(row)

    return normalized


def _is_key_value_table(rows: list[list[str]]) -> bool:
    """
    Detect whether a table is a key-value/definitions table.

    Heuristic: a 2-column table where the second column contains
    significantly longer text than the first column (definitions,
    descriptions, notes) is likely a key-value table where row 0
    is data, not a header.

    Examples that should match:
        | Term           | A roaming agreement is defined as...       |
        | Effective Date | The date on which this agreement becomes... |

    Examples that should NOT match:
        | Rate Type | Amount |
        | MOC       | 0.15   |

    Args:
        rows: Normalized table rows (all same length).

    Returns:
        True if the table matches key-value pattern.
    """
    if not rows:
        return False

    # Only applies to 2-column tables
    if len(rows[0]) != 2:
        return False

    # Need at least 2 rows to make a judgment
    if len(rows) < 2:
        return False

    # Check if second column is consistently longer than the first
    col2_longer_count = 0
    col2_total_length = 0

    for row in rows:
        col1_len = len(row[0])
        col2_len = len(row[1])
        col2_total_length += col2_len

        if col2_len > col1_len:
            col2_longer_count += 1

    col2_avg_length = col2_total_length / len(rows)

    # Both conditions must be true:
    # 1. Second column is longer in most rows
    # 2. Average second-column length exceeds threshold
    ratio = col2_longer_count / len(rows)

    if ratio >= _KV_RATIO_THRESHOLD and col2_avg_length >= _KV_AVG_LENGTH_THRESHOLD:
        logger.debug(
            f"Key-value table detected: ratio={ratio:.2f}, "
            f"avg_col2_len={col2_avg_length:.0f}"
        )
        return True

    # Additional check: if row 0 col 1 looks like a label and col 2
    # looks like a definition (not a short header word), it's key-value.
    # Short headers like "Amount", "Rate", "Zone" are < 20 chars.
    first_row_col2_len = len(rows[0][1])
    if first_row_col2_len > 50 and ratio >= 0.5:
        logger.debug("Key-value table detected via first-row length check")
        return True

    return False
